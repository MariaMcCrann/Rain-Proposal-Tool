"""
Run with:
    python app.py

Then open:
    http://127.0.0.1:5000

.env options:
    TEST_MODE=1   no API calls, uses fixture data
    TEST_MODE=0   real uploaded file workflow
    USE_HAIKU=1   cheaper Claude model for real testing
"""

import os
import uuid
import traceback
from datetime import datetime

from dotenv import load_dotenv
from flask import Flask, request, render_template, send_file
from docx import Document
from openpyxl import Workbook

load_dotenv()

TEST_MODE = os.environ.get("TEST_MODE", "0") == "1"
USE_HAIKU = os.environ.get("USE_HAIKU", "0") == "1"

if USE_HAIKU:
    os.environ["RAIN_MODEL"] = "claude-haiku-4-5"

from ingest import extract_text, ScannedPdfError, UnsupportedFileError
from extract_rfq import extract_rfq
from fill_fee_template import fill_phased_template, rough_fee_estimate
from quotient_url import build_quotient_prefill_url
from draft_proposal_doc import generate_draft_proposal
from generate_proposal import generate_proposal_content
from knowledge_engine.research import run_research

if TEST_MODE:
    from test_fixture import FIXTURE_EXTRACTED, FIXTURE_ESTIMATE


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
TEMPLATE_XLSX = os.path.join(BASE_DIR, "Rain_Project_Fee_Tracker_Template.xlsx")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024


# ---------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------

def get_project_address(extracted: dict, combined_text: str = "") -> str:
    """
    Finds the best available site address from extracted RFQ data.
    Falls back to a simple text search if the extractor did not return one.
    """

    possible_keys = [
        "site_address",
        "project_address",
        "location",
        "site_location",
        "address",
    ]

    for key in possible_keys:
        value = extracted.get(key)
        if value:
            return value

    # Simple fallback for the Avalon style RFQ
    text_upper = combined_text.upper()
    if "15 AVALON ROAD" in text_upper:
        return "15 Avalon Road, Avalon VIC"

    return ""


def create_test_proposal_docx(extracted):
    filename = f"TEST_Draft_Proposal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    doc.add_heading("Draft Proposal", level=1)

    doc.add_heading(extracted.get("project_title", "Untitled Project"), level=2)
    doc.add_paragraph(extracted.get("background", "No background provided."))

    doc.add_heading("Site / Location", level=2)
    doc.add_paragraph(extracted.get("site_address", "Not stated."))

    doc.add_heading("Phases and Deliverables", level=2)
    for phase in extracted.get("phases", []):
        doc.add_heading(phase.get("phase_name", "Unnamed phase"), level=3)
        for item in phase.get("deliverables", []):
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Authority / Standards Requirements", level=2)
    for item in extracted.get("authority_requirements", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.save(path)
    return filename


def create_test_fee_template(extracted):
    filename = f"TEST_Fee_Template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    path = os.path.join(OUTPUT_DIR, filename)

    wb = Workbook()
    ws = wb.active
    ws.title = "Fee Estimate"

    ws["A1"] = "RAIN Consulting - Test Fee Template"
    ws["A3"] = "Project"
    ws["B3"] = extracted.get("project_title", "Untitled Project")

    ws["A5"] = "Phase"
    ws["B5"] = "Hours"
    ws["C5"] = "Fee"

    row = 6
    for phase in extracted.get("phases", []):
        ws[f"A{row}"] = phase.get("phase_name", "Unnamed phase")
        ws[f"B{row}"] = ""
        ws[f"C{row}"] = ""
        row += 1

    wb.save(path)
    return filename


# ---------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------

@app.route("/", methods=["GET"])
def index():
    return render_template(
        "index.html",
        test_mode=TEST_MODE,
        use_haiku=USE_HAIKU,
    )


@app.route("/process", methods=["POST"])
def process():
    mode = request.form.get("mode", "quick")
    include_research = request.form.get("include_research") in ["1", "on", "true"]

    # -------------------------------------------------------------
    # TEST MODE
    # -------------------------------------------------------------
    if TEST_MODE:
        extracted = FIXTURE_EXTRACTED
        warnings = ["⚠️ TEST MODE — this is fixture data, not a real extraction."]

        if mode == "quick":
            return render_template(
                "result.html",
                extracted=extracted,
                warnings=warnings,
                mode="quick",
                source_files=["test_fixture.rfq"],
                estimate_text=None,
                quotient_url=None,
                download_filename=None,
                docx_filename=None,
                include_research=False,
                test_mode=True,
            )

        docx_filename = create_test_proposal_docx(extracted)
        download_filename = create_test_fee_template(extracted)

        return render_template(
            "result.html",
            extracted=extracted,
            warnings=warnings,
            mode="full",
            source_files=["test_fixture.rfq"],
            estimate_text=FIXTURE_ESTIMATE,
            quotient_url="#",
            download_filename=download_filename,
            docx_filename=docx_filename,
            include_research=False,
            test_mode=True,
        )

    # -------------------------------------------------------------
    # NORMAL MODE
    # -------------------------------------------------------------
    uploaded_files = []

    for key in request.files:
    uploaded_files.extend(request.files.getlist(key))

    uploaded_files = [f for f in uploaded_files if f and f.filename]


    if not uploaded_files:
        return render_template(
            "index.html",
            error="No file selected.",
            test_mode=TEST_MODE,
            use_haiku=USE_HAIKU,
        )

    run_id = uuid.uuid4().hex[:8]
    text_sections = []
    skipped = []

    for uploaded in uploaded_files:
        safe_name = uploaded.filename.replace("/", "_").replace("\\", "_")
        upload_path = os.path.join(UPLOAD_DIR, f"{run_id}_{safe_name}")
        uploaded.save(upload_path)

        try:
            file_text = extract_text(upload_path)
            text_sections.append(f"--- Source file: {uploaded.filename} ---\n{file_text}")
        except ScannedPdfError as e:
            skipped.append(f"{uploaded.filename}: {e}")
        except UnsupportedFileError as e:
            skipped.append(f"{uploaded.filename}: {e}")
        except Exception:
            skipped.append(
                f"{uploaded.filename}: couldn't read "
                f"({traceback.format_exc(limit=1).strip()})"
            )

    if not text_sections:
        return render_template(
            "index.html",
            error="None of the uploaded files could be read:\n" + "\n".join(skipped),
            test_mode=TEST_MODE,
            use_haiku=USE_HAIKU,
        )

    combined_text = "\n\n".join(text_sections)

    try:
        extracted = extract_rfq(combined_text)
    except Exception as e:
        return render_template(
            "index.html",
            error=f"Extraction failed: {e}",
            test_mode=TEST_MODE,
            use_haiku=USE_HAIKU,
        )

    warnings = list(skipped)

    # -------------------------------------------------------------
    # KNOWLEDGE ENGINE
    # -------------------------------------------------------------
    research = None
    project_address = get_project_address(extracted, combined_text)

    if include_research:
        if not project_address:
            warnings.append("Knowledge Engine skipped: no project address found.")
        else:
            try:
                research = run_research(project_address)

                print("\n=== KNOWLEDGE ENGINE RESULT ===")
                print(research)
                print("================================\n")

                for note in research.get("notes", []):
                    warnings.append(f"Research note: {note}")

            except Exception as e:
                warnings.append(f"Knowledge Engine failed: {e}")

    if mode == "quick":
        return render_template(
            "result.html",
            extracted=extracted,
            warnings=warnings,
            mode="quick",
            source_files=[f.filename for f in uploaded_files],
            estimate_text=None,
            quotient_url=None,
            download_filename=None,
            docx_filename=None,
            include_research=include_research,
            test_mode=False,
        )

    # -------------------------------------------------------------
    # FULL MODE
    # -------------------------------------------------------------
    output_filename = f"{run_id}_fee_template.xlsx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    try:
        fill_warnings = fill_phased_template(TEMPLATE_XLSX, output_path, extracted)
        warnings += fill_warnings
    except Exception as e:
        warnings.append(f"Fee template failed to generate: {e}")
        output_filename = None

    try:
        sections = generate_proposal_content(extracted, research)
    except Exception as e:
        warnings.append(f"Proposal content generation failed: {e}")
        sections = {}

    docx_filename = f"{run_id}_draft_proposal.docx"
    docx_path = os.path.join(OUTPUT_DIR, docx_filename)

    try:
        generate_draft_proposal(extracted, sections, docx_path, research)
    except Exception as e:
        warnings.append(f"Draft proposal document failed to generate: {e}")
        docx_filename = None

    try:
        estimate_text = rough_fee_estimate(extracted)
    except Exception:
        estimate_text = None

    try:
        quotient_url = build_quotient_prefill_url(extracted)
    except Exception:
        quotient_url = None

    return render_template(
        "result.html",
        extracted=extracted,
        warnings=warnings,
        estimate_text=estimate_text,
        quotient_url=quotient_url,
        download_filename=output_filename,
        docx_filename=docx_filename,
        source_files=[f.filename for f in uploaded_files],
        include_research=include_research,
        mode="full",
        test_mode=False,
    )


@app.route("/download/<filename>")
def download(filename):
    if not filename or filename == "None":
        return "No file was generated for download.", 404

    path = os.path.join(OUTPUT_DIR, filename)

    if not os.path.exists(path):
        return f"File not found: {filename}", 404

    return send_file(path, as_attachment=True, download_name=filename)


if __name__ == "__main__":
    if TEST_MODE:
        print(">>> TEST MODE ON — no API calls will be made.")
    elif USE_HAIKU:
        print(">>> USE_HAIKU ON — using claude-haiku-4-5.")

    if not TEST_MODE and not os.environ.get("ANTHROPIC_API_KEY"):
        print("WARNING: ANTHROPIC_API_KEY is not set — extraction will fail.")

    app.run(debug=True, port=5000)