# draft_proposal_doc.py

from typing import Any, Dict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraphs(doc: Document, text: str):
    """
    Adds text as proper Word paragraphs.
    Splits on blank lines so the Project Understanding does not appear as one ugly block.
    """

    if not text:
        return

    for para in text.split("\n\n"):
        cleaned = para.strip()
        if cleaned:
            p = doc.add_paragraph(cleaned)
            p.paragraph_format.space_after = Pt(8)


def add_site_research_table(doc: Document, research: Dict[str, Any]):
    rows = [
        ("Site address", research.get("site_address", "Not identified")),
        ("Latitude", research.get("latitude", "Not identified")),
        ("Longitude", research.get("longitude", "Not identified")),
        ("Council", research.get("council", "Not identified")),
        ("Traditional Owners", research.get("traditional_owners", "Not identified")),
        ("CMA", research.get("cma", "Not identified")),
        ("Water authority", research.get("water_authority", "Not identified")),
        ("Zone", research.get("zone", "To be confirmed")),
        ("DPO", research.get("dpo", "To be confirmed")),
        ("SBO", research.get("sbo", "To be confirmed")),
        ("LSIO", research.get("lsio", "To be confirmed")),
        ("FO", research.get("fo", "To be confirmed")),
        ("Planning source", research.get("planning_source", "VicPlan / planning research")),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Preliminary finding"

    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = str(value)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def apply_document_style(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def create_draft_proposal_doc(
    output_path: str,
    extracted: Dict[str, Any],
    research: Dict[str, Any],
    sections: Dict[str, Any],
):
    doc = Document()
    apply_document_style(doc)

    project_title = extracted.get("project_title", "Draft Proposal")
    site_address = research.get("site_address") or extracted.get("site_address", "")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAIN Consulting")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Draft Proposal")
    run.bold = True
    run.font.size = Pt(14)

    if site_address:
        address_para = doc.add_paragraph()
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_para.add_run(site_address)

    doc.add_paragraph("")

    add_heading(doc, "1. Project Understanding", level=1)
    add_paragraphs(doc, sections.get("project_understanding", ""))

    add_heading(doc, "2. Scope of Services", level=1)
    add_paragraphs(doc, sections.get("scope_of_services", ""))

    add_heading(doc, "3. Preliminary Site Research", level=1)
    intro = (
        "The following preliminary site information has been identified to assist with "
        "scoping and proposal preparation. Planning controls and authority requirements "
        "should be confirmed through VicPlan and the relevant authorities before finalising "
        "the technical approach."
    )
    add_paragraphs(doc, intro)

    add_site_research_table(doc, research)

    add_heading(doc, "4. Assumptions and Exclusions", level=1)
    assumptions = (
        "This proposal has been prepared based on the information provided at the time of "
        "preparation. The scope may need to be reviewed if additional authority requirements, "
        "planning controls, survey information, development layouts, or modelling requirements "
        "are identified.\n\n"
        "Unless specifically included in the final agreed scope, detailed hydraulic modelling, "
        "civil design documentation, authority application fees, survey, geotechnical inputs, "
        "environmental assessments, and legal advice are excluded."
    )
    add_paragraphs(doc, assumptions)

    doc.save(output_path)

    return output_path