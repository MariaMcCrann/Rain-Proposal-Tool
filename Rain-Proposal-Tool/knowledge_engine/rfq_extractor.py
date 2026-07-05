# extract_rfq.py

import os
import re
from typing import Dict, Any

from pypdf import PdfReader
from docx import Document

from research_site import extract_site_address


def read_pdf_text(file_path: str) -> str:
    text_parts = []

    reader = PdfReader(file_path)

    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_parts.append(page_text)

    return "\n".join(text_parts).strip()


def read_docx_text(file_path: str) -> str:
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    return "\n".join(text_parts).strip()


def read_uploaded_file_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return read_pdf_text(file_path)

    if ext == ".docx":
        return read_docx_text(file_path)

    if ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    raise ValueError(f"Unsupported file type: {ext}")


def clean_project_title(filename: str) -> str:
    name = os.path.basename(filename)
    name = os.path.splitext(name)[0]
    name = name.replace("_", " ").replace("-", " ")
    name = re.sub(r"\s+", " ", name).strip()
    return name


def extract_project_type(text: str) -> str:
    text_l = text.lower()

    if "subdivision" in text_l:
        return "Subdivision"
    if "industrial" in text_l:
        return "Industrial development"
    if "residential" in text_l:
        return "Residential development"
    if "stormwater" in text_l:
        return "Stormwater assessment"
    if "flood" in text_l:
        return "Flood assessment"
    if "hydrological" in text_l or "hydrology" in text_l:
        return "Hydrological engineering services"

    return "Engineering services"


def extract_scope_summary(text: str) -> str:
    """
    Keeps this simple and stable.
    The proposal writer can expand this later.
    """

    text_l = text.lower()

    if "phase 3" in text_l or "phase 4" in text_l or "phase 5" in text_l:
        return "Hydrological engineering services for remaining project phases."

    if "scope of works" in text_l:
        return "Engineering scope of works based on the RFQ."

    return "Engineering services based on the RFQ."


def extract_rfq(file_path: str) -> Dict[str, Any]:
    """
    Main RFQ extractor.

    Returns:
    - full_text
    - project_title
    - site_address
    - project_type
    - scope_summary
    """

    full_text = read_uploaded_file_text(file_path)
    project_title = clean_project_title(file_path)

    site_address = extract_site_address(full_text, project_title)

    extracted = {
        "file_name": os.path.basename(file_path),
        "project_title": project_title,
        "site_address": site_address,
        "project_type": extract_project_type(full_text),
        "scope_summary": extract_scope_summary(full_text),
        "full_text": full_text,
    }

    return extracted

def _derive_project_title(full_text: str) -> str:
    if not full_text:
        return "Untitled Project"

    patterns = [
        r"(?:project\s*title|project\s*name|project)\s*[:\-]\s*(.+)",
        r"(?:re|subject)\s*[:\-]\s*(.+)",
    ]

    for pattern in patterns:
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            title = re.split(r"[\r\n]", title)[0].strip()
            if title:
                return title

    for line in full_text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("--- Source file:") and line.endswith("---"):
            continue
        return line[:120]

    return "Untitled Project"

def extract_rfq_data(full_text: str) -> Dict[str, Any]:
    """
    Extracts RFQ fields directly from already-extracted text (e.g. text
    already read from one or more uploaded files and combined by the
    caller). Unlike extract_rfq(), this does not read files from disk.

    Returns:
    - full_text
    - project_title
    - site_address
    - project_type
    - scope_summary
    """
    project_title = _derive_project_title(full_text)
    site_address = extract_site_address(full_text, project_title)

    return {
        "project_title": project_title,
        "site_address": site_address,
        "project_type": extract_project_type(full_text),
        "scope_summary": extract_scope_summary(full_text),
        "background": extract_background(full_text),
        "phases": extract_phases(full_text),
        "authority_requirements": extract_authority_requirements(full_text),
        "contact": extract_contact(full_text),
        "extraction_notes": [],
        "full_text": full_text,
    }
def extract_background(full_text: str) -> str:
    if not full_text:
        return ""

    patterns = [
        r"(?:background|overview|project\s*background)\s*[:\-]\s*(.+?)(?:\n\s*\n|\Z)",
    ]

    for pattern in patterns:
        match = re.search(pattern, full_text, re.IGNORECASE | re.DOTALL)
        if match:
            value = re.sub(r"\s+", " ", match.group(1)).strip()
            if value:
                return value

    return ""


def extract_authority_requirements(full_text: str) -> list:
    if not full_text:
        return []

    keywords = [
        "guideline", "standard", "planning scheme", "regulation",
        "code of practice", "australian standard", "policy",
    ]

    results = []
    for line in full_text.splitlines():
        line_stripped = line.strip(" -\u2022\t")
        if not line_stripped:
            continue
        line_l = line_stripped.lower()
        if any(keyword in line_l for keyword in keywords):
            if line_stripped not in results:
                results.append(line_stripped)

    return results[:10]

def extract_contact(full_text: str) -> Dict[str, str]:
    contact = {"name": "", "email": "", "phone": "", "company": ""}

    if not full_text:
        return contact

    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", full_text)
    if email_match:
        contact["email"] = email_match.group(0)

    phone_match = re.search(r"(?:\+?\d[\d ()\-]{7,}\d)", full_text)
    if phone_match:
        contact["phone"] = phone_match.group(0).strip()

    name_match = re.search(r"(?:contact\s*(?:name|person)?|attention)\s*[:\-]\s*(.+)", full_text, re.IGNORECASE)
    if name_match:
        contact["name"] = re.split(r"[\r\n]", name_match.group(1).strip())[0].strip()

    company_match = re.search(r"(?:company|organisation|organization|council)\s*[:\-]\s*(.+)", full_text, re.IGNORECASE)
    if company_match:
        contact["company"] = re.split(r"[\r\n]", company_match.group(1).strip())[0].strip()

    return contact

def extract_phases(full_text: str) -> list:
    if not full_text:
        return []

    phase_pattern = re.compile(r"phase\s*\d+[^\n]*", re.IGNORECASE)
    matches = list(phase_pattern.finditer(full_text))
    phases = []

    for i, match in enumerate(matches):
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        phase_name = match.group(0).strip(" -:\u2022\t")
        body = full_text[start:end]

        deliverables = []
        for line in body.splitlines():
            line_stripped = line.strip(" -\u2022\t")
            if line_stripped:
                deliverables.append(line_stripped)
            if len(deliverables) >= 8:
                break

        phases.append({"phase_name": phase_name, "deliverables": deliverables})

    return phases
